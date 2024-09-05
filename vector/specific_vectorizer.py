from docx import Document as DocxDocument
from langchain_community.document_loaders import UnstructuredWordDocumentLoader	
import pandas as pd
from collections import defaultdict
import re
import os

from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_text_splitters import CharacterTextSplitter
from langchain.schema.document import Document
from langchain_community.vectorstores import Chroma

from typing import Optional
from typing import List
from pathlib import Path
import shutil

# insert Google Gemini API Key
os.environ['GOOGLE_API_KEY'] = 'insert-api-key'

# Path to the docx file
docx_path = "/data/documents/sample.docx"  # Replace with the path to your docx file
abs_path = Path().absolute()
directory = f'{abs_path}/data/documents/'
document_paths = [os.path.join(directory, filename) for filename in os.listdir(directory) if filename.endswith('.docx')]
folder_path = f'{abs_path}/database'

def check_and_delete_folder(folder_path):
    if os.path.exists(folder_path):
        print(f"Folder '{folder_path}' exists. Deleting it.")
        shutil.rmtree(folder_path)
        print(f"Folder '{folder_path}' has been deleted.")
    else:
        return

check_and_delete_folder(folder_path)

titles = []
headers = []
sections = []
documents = []

class CustomGoogleGenerativeAIEmbeddings(GoogleGenerativeAIEmbeddings):
    def embed_documents(self, texts: List[str], task_type: Optional[str] = None, titles: Optional[List[str]] = None, output_dimensionality: Optional[int] = None) -> List[List[float]]:
        embeddings_repeated = super().embed_documents(texts, task_type, titles, output_dimensionality)
        # convert proto.marshal.collections.repeated.Repeated to list
        embeddings = [list(emb) for emb in embeddings_repeated]
        return embeddings

gemini_embeddings_wrapper = CustomGoogleGenerativeAIEmbeddings(model="models/embedding-001")

def vectorize(document_paths):
    '''
    Vectorize documents stored in document path specified based on specific document format
    
    For this use case: document is split into five sections with different headings, hence splitting/chunking is done by sections with the Incident title appended to each section to identify which incident each section belongs to
    
    If specific format is not necessary, basic token count chunking with overlap can also be used
    '''
    
    documents = []
    for filename in document_paths: 
        document_loader = UnstructuredWordDocumentLoader(filename)
        document = document_loader.load()
        title = re.search(r"INCIDENT TITLE:\s*(.*?)\s*INCIDENT SEVERITY:", document[0].page_content)
        clean_title = title.group().removeprefix("INCIDENT TITLE: ").removesuffix("INCIDENT SEVERITY:")

        rx = re.compile(r"(SECTION).*", re.VERBOSE | re.MULTILINE)

        for match in rx.finditer(document[0].page_content):
            if "SECTION I: INCIDENT INFORMATION " in match.group():
                section = match.group().removeprefix("SECTION I: INCIDENT INFORMATION ")
                metadata = {'incident_title': clean_title, 'section_title':"SECTION I: INCIDENT INFORMATION"}
                new_section = 'INCIDENT TITLE: ' + clean_title + '\n' + section
                chunk = Document(page_content=new_section, metadata=metadata)
                documents += [chunk]

                sections.append(section)
                titles.append(clean_title)
                headers.append("SECTION I: INCIDENT INFORMATION")

            if "SECTION II: INCIDENT INFORMATION " in match.group():
                section = match.group().removeprefix("SECTION II: INCIDENT INFORMATION ")
                metadata = {'incident_title': clean_title, 'section_title':"SECTION II: INCIDENT INFORMATION"}
                new_section = 'INCIDENT TITLE: ' + clean_title + '\n' + section
                chunk = Document(page_content=new_section, metadata=metadata)
                documents += [chunk]

                sections.append(section)
                titles.append(clean_title)
                headers.append("SECTION II: INCIDENT INFORMATION")

            if "SECTION III: 5-WHY ANALYSIS PROCESS " in match.group():
                section = match.group().removeprefix("SECTION III: 5-WHY ANALYSIS PROCESS ")
                metadata = {'incident_title': clean_title, 'section_title':"SECTION III: 5-WHY ANALYSIS PROCESS"}
                new_section = 'INCIDENT TITLE: ' + clean_title + '\n' + 'Causal Factors and Root Cause Analysis \n' + section
                chunk = Document(page_content=new_section, metadata=metadata)
                documents += [chunk]

                sections.append(section)
                titles.append(clean_title)
                headers.append("SECTION III: 5-WHY ANALYSIS PROCESS")
            
            if "SECTION IV: MEASURES FOR PREVENTION OF RECURRENCE " in match.group():
                section = match.group().removeprefix("SECTION IV: MEASURES FOR PREVENTION OF RECURRENCE RECOMMENDATION AND IMPLEMENTATION OF MEASURES TO ADDRESS ROOT CAUSE AND CAUSAL FACTORS. (Please including corrective actions taken by Agency, e.g. counselling, warning or training the staff/vendor involved,  enhanced Agency’s incident prevention awareness programme, and improved Agency’s ICT security and data governance, plan, procedures and processes, controls and reviews):")
                metadata = {'incident_title': clean_title, 'section_title':"SECTION IV: MEASURES FOR PREVENTION OF RECURRENCE"}
                new_section = 'INCIDENT TITLE: ' + clean_title + '\n' + 'Measures for Prevention \n' + section
                chunk = Document(page_content=new_section, metadata=metadata)
                documents += [chunk]

                sections.append(section)
                titles.append(clean_title)
                headers.append("SECTION IV: MEASURES FOR PREVENTION OF RECURRENCE")

            if "SECTION IV: REVIEW OF AGENCY’S INCIDENT MANAGEMENT READINESS " in match.group():
                section = match.group().removeprefix("SECTION IV: REVIEW OF AGENCY’S INCIDENT MANAGEMENT READINESS ")
                metadata = {'incident_title': clean_title, 'section_title':"SECTION IV: REVIEW OF AGENCY’S INCIDENT MANAGEMENT READINESS"}
                new_section = 'INCIDENT TITLE: ' + clean_title + '\n' + 'Review of Incident Management \n' + section
                chunk = Document(page_content=new_section, metadata=metadata)
                documents += [chunk]

                sections.append(section)
                titles.append(clean_title)
                headers.append("SECTION IV: REVIEW OF AGENCY’S INCIDENT MANAGEMENT READINESS")
    
    print(f"Documents from '{document_paths}' loaded.")
    vectorstore = Chroma.from_documents(
                    documents=documents,  # Data
                    embedding=gemini_embeddings_wrapper,    # Embedding model
                    persist_directory="./database" # Directory to save data
                    )
    print(f"Vector database {abs_path}/database created.")

vectorize(document_paths)

# to store dataframe if necessary
#df = pd.DataFrame()
#df['incident_title'] = titles
#df['section_title'] = headers
#df['content'] = sections


# Convert the extracted data to a DataFrame
#df = pd.DataFrame(sections)

# Save the DataFrame to an Excel file
#output_path = "/data/incident_samples.xlsx"  # Replace with your desired output path
#df.to_excel(output_path, index=False)

#print(f"Extracted sections have been saved to {output_path}")
